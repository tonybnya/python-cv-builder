from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture(
    'me.jpg', width=Inches(1.0)
)

print('')
print('-------------------------------------------------')
print("-   A program which helps you build your CV   -")
print('-------------------------------------------------')
print('')
print('Let me ask some details about you...')
print('')

# name, phone and email details
name = input("What's your name? -> ")
#speak('Hello ' + name + '!' + ' How are you today?')
phone = input("What's your phone number? -> ")
email = input("What's your email? -> ")

print('')

# Set an paragraph with name, phone and email details
document.add_paragraph(
    name + ' | ' + phone + ' | ' + email)

print('')
print('-------------------------------------------------')
print('------------      ABOUT YOU        ------------')
print('-------------------------------------------------')
print('')

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell something about yourself... -> ')
)

print('')
print('------------------------------------------------')
print('---      Skip to your Work Experiences     ---')
print('------------------------------------------------')
print('')

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter a company name -> ')
fromDate = input('From Date -> ')
toDate = input('To Date -> ')
print('')

p.add_run(company + ' ->' + ' ').bold = True
p.add_run('('+fromDate + ' - ' + toDate + ')' + '\n').italic = True

experienceDetails = input(
    'Describe your experience at ' + company + ' -> '
)
p.add_run(experienceDetails)

# more experiences
while True:
    hasMoreExperiences = input(
        'Do you have more experiences? (yes = YES ; no = NO) -> '
    )

    print('')

    if hasMoreExperiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter a company name -> ')
        fromDate = input('From Date -> ')
        toDate = input('To Date -> ')
        print('')

        p.add_run(company + ' ->' + ' ').bold = True
        p.add_run('('+fromDate + ' - ' + toDate + ')' + '\n').italic = True

        experienceDetails = input(
            'Describe your experience at ' + company + ': ' + ' -> '
        )

        p.add_run(experienceDetails)
    else:
        break

print('------------------------------------------------')
print('--------     Skip to your Skills     ---------')
print('------------------------------------------------')
print('')

# skills
document.add_heading('Skills')
skill = input('Enter a skill -> ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'
print('')

while True:
    hasMoreSkills = input(
        'Do you have more skills? (yes = YES ; no = NO) -> '
    )
    print('')

    if hasMoreSkills.lower() == 'yes':
        skill = input('Enter another skill -> ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

print('------------------------------------------------')
print('------     Skip to your References     -------')
print('------------------------------------------------')
print('')

# references
document.add_heading('References')

ref = []

person = input('Enter a person to contact about you -> ')
r_email = input('Enter his email address -> ')
r_phone = input('Enter his phone number -> ')
r_company = input('Enter his company ->')
print('')

ref.append(person)
ref.append(r_email)
ref.append(r_phone)
ref.append(r_company)

for item in range(len(ref)):
    p = document.add_paragraph(item)

person.style = 'List Bullet'

while True:
    hasMoreReferences = input(
        'Do you have more references? (yes = YES ; no = NO) -> '
    )
    print('')

    if hasMoreReferences == 'yes':
        ref = []

        person = input('Enter a person to contact about you -> ')
        r_email = input('Enter his email address -> ')
        r_phone = input('Enter his phone number -> ')
        r_company = input('Enter his company ->')
        print('')

        ref.append(person)
        ref.append(r_email)
        ref.append(r_phone)
        ref.append(r_company)

        for item in range(len(ref)):
            p = document.add_paragraph(item)

        person.style = 'List Bullet'
    else:
        break

print('------------------------------------------------')
print('----------     Footer goes down     ----------')
print('------------------------------------------------')
print('')

footerYear = input('Enter the current year -> ')
footerName = input('Enter the builder name to print at the footer -> ')
print('')
print('The footer of your CV will appear like this down...')
print('')
print(
    '© Copyright ' + footerYear + ' by ' +
    footerName + ' - CV generated using Python.'
)

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "© Copyright 2020, by Tony B. NYA - CV generated using Python."

document.save('cv.docx')
